from sys import path
from numpy.lib.function_base import select
import pandas as pd
from docx import Document
from docx.shared import Inches


class WordTesting:
    def __init__(self):
        # initialize extra class.
        self.dataQuestion = self.dataAnswer = ['']  # 初始化问题和答案
        self.count = 3  # 生成文件的数量
        self.ReadExcel()
        self.CreateFile()

    def ReadExcel(self):
        # 读取数据,设置None可以生成一个字典，字典中的key值即为sheet名字，此时不用使用DataFram，会报错
        path = "data.xls"
        data = pd.DataFrame(pd.read_excel(path))
        self.dataQuestion = data['问题']  # 获取列名为问题这一列的内容
        self.dataAnswer = data['答案']  # 获取列名为答案这一列的内容
    # 生成序列文件

    def CreateFile(self):
        # 当需要生成的文件数量大于0时，继续循环，直到全部试卷生成完成为止
        while self.count > 0:
            self.count -= 1
            document = Document()
            document.add_heading("考试试卷", 0)
            for i in range(0, len(self.dataQuestion)):
                # 添加自然段 ，如果答案为空则生成填空题
                if str(self.dataAnswer[i]) == "nan":
                    document.add_paragraph(
                        str(self.dataQuestion[i])+"\n" + "_"*300+"\n", style='List Number'
                    )
                # 添加自然段 ，如果答案为不为空则根据答案类型生成判断或者选择题
                else:

                    document.add_paragraph(
                        str(self.dataQuestion[i])+"\n" + str(self.dataAnswer[i]) + "\n", style='List Number'
                    )
            # 添加换行符
            document.add_page_break()
            document.save("data/" + str(self.count)+".docx")


if __name__ == "__main__":
    WordTesting()
